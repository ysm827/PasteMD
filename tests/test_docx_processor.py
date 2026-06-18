import io
import re
import zipfile
from xml.etree import ElementTree as ET

from docx import Document

from pastemd.config.defaults import DEFAULT_CONFIG
from pastemd.utils.docx_processor import DocxProcessor


NAMESPACES = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}


def _make_label_content_docx() -> bytes:
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    rows = [
        ("要素", "内容"),
        ("相关知识点\n(讲义原文)", "中位数：把一组数据按从小到大或大到小进行排列，位置居中的数值叫做中位数。"),
        ("题目", "2019 年某县有 8 个大型建筑企业，这 8 个企业的员工人数分别为 600、632、710、740、760、800、850、910。"),
    ]
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def _make_docx_table(rows: list[tuple[str, ...]]) -> bytes:
    doc = Document()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def _table_grid_widths(docx_bytes: bytes) -> list[int]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        document_xml = archive.read("word/document.xml")

    root = ET.fromstring(document_xml)
    grid_cols = root.findall(".//w:tbl[1]/w:tblGrid/w:gridCol", NAMESPACES)
    return [
        int(col.get(f"{{{NAMESPACES['w']}}}w"))
        for col in grid_cols
        if col.get(f"{{{NAMESPACES['w']}}}w")
    ]


def _document_xml(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        return archive.read("word/document.xml").decode("utf-8")


def _table_width_attrs(docx_bytes: bytes) -> tuple[str | None, str | None]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        document_xml = archive.read("word/document.xml")

    root = ET.fromstring(document_xml)
    table_width = root.find(".//w:tbl[1]/w:tblPr/w:tblW", NAMESPACES)
    assert table_width is not None
    return (
        table_width.get(f"{{{NAMESPACES['w']}}}type"),
        table_width.get(f"{{{NAMESPACES['w']}}}w"),
    )


def _table_property_child_names(docx_bytes: bytes) -> list[str]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        document_xml = archive.read("word/document.xml")

    root = ET.fromstring(document_xml)
    table_properties = root.find(".//w:tbl[1]/w:tblPr", NAMESPACES)
    assert table_properties is not None
    return [child.tag.rsplit("}", 1)[-1] for child in list(table_properties)]


def _first_cell_property_child_names(docx_bytes: bytes) -> list[str]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        document_xml = archive.read("word/document.xml")

    root = ET.fromstring(document_xml)
    cell_properties = root.find(".//w:tbl[1]/w:tr[1]/w:tc[1]/w:tcPr", NAMESPACES)
    assert cell_properties is not None
    return [child.tag.rsplit("}", 1)[-1] for child in list(cell_properties)]


def _set_first_table_width(docx_bytes: bytes, width: int) -> bytes:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        document_xml = archive.read("word/document.xml")
        root = ET.fromstring(document_xml)
        table = root.find(".//w:tbl[1]", NAMESPACES)
        assert table is not None

        table_properties = table.find("./w:tblPr", NAMESPACES)
        if table_properties is None:
            table_properties = ET.Element(f"{{{NAMESPACES['w']}}}tblPr")
            table.insert(0, table_properties)

        table_width = table_properties.find("./w:tblW", NAMESPACES)
        if table_width is None:
            table_width = ET.SubElement(table_properties, f"{{{NAMESPACES['w']}}}tblW")
        table_width.set(f"{{{NAMESPACES['w']}}}type", "dxa")
        table_width.set(f"{{{NAMESPACES['w']}}}w", str(width))

        grid_columns = table.findall("./w:tblGrid/w:gridCol", NAMESPACES)
        assert grid_columns
        base_width = width // len(grid_columns)
        for index, column in enumerate(grid_columns):
            column_width = base_width if index < len(grid_columns) - 1 else width - base_width * index
            column.set(f"{{{NAMESPACES['w']}}}w", str(column_width))

        for row in table.findall("./w:tr", NAMESPACES):
            cells = row.findall("./w:tc", NAMESPACES)
            for index, cell in enumerate(cells):
                cell_properties = cell.find("./w:tcPr", NAMESPACES)
                if cell_properties is None:
                    cell_properties = ET.Element(f"{{{NAMESPACES['w']}}}tcPr")
                    cell.insert(0, cell_properties)
                cell_width = cell_properties.find("./w:tcW", NAMESPACES)
                if cell_width is None:
                    cell_width = ET.SubElement(cell_properties, f"{{{NAMESPACES['w']}}}tcW")
                column_width = base_width if index < len(grid_columns) - 1 else width - base_width * index
                cell_width.set(f"{{{NAMESPACES['w']}}}type", "dxa")
                cell_width.set(f"{{{NAMESPACES['w']}}}w", str(column_width))

        updated_document_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        output_stream = io.BytesIO()
        with zipfile.ZipFile(output_stream, "w") as output:
            for item in archive.infolist():
                data = updated_document_xml if item.filename == "word/document.xml" else archive.read(item.filename)
                output.writestr(item, data)
        return output_stream.getvalue()


def _set_first_table_percentage_width(docx_bytes: bytes, width: int) -> bytes:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        document_xml = archive.read("word/document.xml")
        root = ET.fromstring(document_xml)
        table_width = root.find(".//w:tbl[1]/w:tblPr/w:tblW", NAMESPACES)
        assert table_width is not None
        table_width.set(f"{{{NAMESPACES['w']}}}type", "pct")
        table_width.set(f"{{{NAMESPACES['w']}}}w", str(width))

        updated_document_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        output_stream = io.BytesIO()
        with zipfile.ZipFile(output_stream, "w") as output:
            for item in archive.infolist():
                data = updated_document_xml if item.filename == "word/document.xml" else archive.read(item.filename)
                output.writestr(item, data)
        return output_stream.getvalue()


def _remove_first_table_width(docx_bytes: bytes) -> bytes:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        document_xml = archive.read("word/document.xml")
        root = ET.fromstring(document_xml)
        table_properties = root.find(".//w:tbl[1]/w:tblPr", NAMESPACES)
        assert table_properties is not None
        table_width = table_properties.find("./w:tblW", NAMESPACES)
        assert table_width is not None
        table_properties.remove(table_width)

        updated_document_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        output_stream = io.BytesIO()
        with zipfile.ZipFile(output_stream, "w") as output:
            for item in archive.infolist():
                data = updated_document_xml if item.filename == "word/document.xml" else archive.read(item.filename)
                output.writestr(item, data)
        return output_stream.getvalue()


def _remove_first_cell_width_and_add_shading(docx_bytes: bytes) -> bytes:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        document_xml = archive.read("word/document.xml")
        root = ET.fromstring(document_xml)
        cell_properties = root.find(".//w:tbl[1]/w:tr[1]/w:tc[1]/w:tcPr", NAMESPACES)
        assert cell_properties is not None
        cell_width = cell_properties.find("./w:tcW", NAMESPACES)
        assert cell_width is not None
        cell_properties.remove(cell_width)
        shading = ET.SubElement(cell_properties, f"{{{NAMESPACES['w']}}}shd")
        shading.set(f"{{{NAMESPACES['w']}}}fill", "FFFF00")

        updated_document_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        output_stream = io.BytesIO()
        with zipfile.ZipFile(output_stream, "w") as output:
            for item in archive.infolist():
                data = updated_document_xml if item.filename == "word/document.xml" else archive.read(item.filename)
                output.writestr(item, data)
        return output_stream.getvalue()


def test_docx_auto_table_layout_is_disabled_by_default():
    assert DEFAULT_CONFIG["docx_auto_table_layout"] is False


def test_apply_custom_processing_keeps_table_widths_when_auto_layout_disabled():
    docx_bytes = _make_label_content_docx()

    processed = DocxProcessor.apply_custom_processing(
        docx_bytes,
        auto_layout_tables=False,
    )

    assert _table_grid_widths(processed) == _table_grid_widths(docx_bytes)


def test_apply_custom_processing_can_auto_layout_label_content_tables():
    docx_bytes = _make_label_content_docx()

    processed = DocxProcessor.apply_custom_processing(
        docx_bytes,
        auto_layout_tables=True,
    )

    first_col, second_col = _table_grid_widths(processed)
    assert first_col < second_col
    assert first_col <= 2200


def test_auto_layout_preserves_existing_fixed_table_width():
    docx_bytes = _set_first_table_width(_make_label_content_docx(), 4000)

    processed = DocxProcessor.apply_custom_processing(
        docx_bytes,
        auto_layout_tables=True,
    )

    assert sum(_table_grid_widths(processed)) == 4000


def test_auto_layout_preserves_existing_percentage_table_width():
    docx_bytes = _set_first_table_percentage_width(_make_label_content_docx(), 2500)

    processed = DocxProcessor.apply_custom_processing(
        docx_bytes,
        auto_layout_tables=True,
    )

    assert _table_width_attrs(processed) == ("pct", "2500")
    assert sum(_table_grid_widths(processed)) == 4680


def test_auto_layout_inserts_table_layout_before_table_look():
    docx_bytes = _make_label_content_docx()

    processed = DocxProcessor.apply_custom_processing(
        docx_bytes,
        auto_layout_tables=True,
    )

    child_names = _table_property_child_names(processed)
    assert child_names.index("tblLayout") < child_names.index("tblLook")


def test_auto_layout_inserts_missing_table_width_before_table_look():
    docx_bytes = _remove_first_table_width(_make_label_content_docx())

    processed = DocxProcessor.apply_custom_processing(
        docx_bytes,
        auto_layout_tables=True,
    )

    child_names = _table_property_child_names(processed)
    assert child_names.index("tblW") < child_names.index("tblLook")
    assert child_names.index("tblLayout") < child_names.index("tblLook")


def test_auto_layout_inserts_missing_cell_width_before_cell_shading():
    docx_bytes = _remove_first_cell_width_and_add_shading(_make_label_content_docx())

    processed = DocxProcessor.apply_custom_processing(
        docx_bytes,
        auto_layout_tables=True,
    )

    child_names = _first_cell_property_child_names(processed)
    assert child_names.index("tcW") < child_names.index("shd")


def test_auto_layout_keeps_wide_table_column_widths_positive():
    widths = DocxProcessor._column_widths_from_scores([1.0] * 12, 9360)

    assert len(widths) == 12
    assert sum(widths) == 9360
    assert all(width > 0 for width in widths)


def test_auto_layout_preserves_ignorable_namespace_declarations():
    docx_bytes = _make_label_content_docx()

    processed = DocxProcessor.apply_custom_processing(
        docx_bytes,
        auto_layout_tables=True,
    )

    xml = _document_xml(processed)
    document_start_match = re.search(r"<w:document\b[^>]*>", xml)
    assert document_start_match is not None
    document_start = document_start_match.group(0)
    ignorable_match = re.search(r"(?:mc|ns\d+):Ignorable=\"([^\"]+)\"", document_start)

    assert ignorable_match is not None
    for prefix in ignorable_match.group(1).split():
        assert f"xmlns:{prefix}=" in document_start


def test_auto_layout_keeps_quiz_explanation_column_readable():
    docx_bytes = _make_docx_table([
        ("模块", "题型", "题干", "选项", "答案", "解析"),
        (
            "描述统计",
            "单选题",
            "某班 8 名学生的成绩分别为 60、62、71、74、76、80、85、91，这组数据的中位数是（ ）分。",
            "A. 74 B. 76 C. 75 D. 80",
            "C",
            "数据已经从小到大排列，数量为 8，是偶数，因此中位数为第 4 个数和第 5 个数的平均值，即 (74+76)/2=75。",
        ),
        (
            "概率",
            "单选题",
            "一个袋子中有 3 个红球、2 个白球、5 个黑球，随机摸出一个球，摸到红球的概率是（ ）。",
            "A. 1/5 B. 3/10 C. 1/2 D. 3/5",
            "B",
            "总球数为 10，红球数为 3，因此概率为 3/10。",
        ),
        (
            "几何",
            "多选题",
            "下列说法正确的是（ ）。",
            "A. 正方形是矩形 B. 矩形一定是正方形 C. 平行四边形对边相等 D. 菱形四条边相等",
            "ACD",
            "正方形属于特殊矩形；矩形不一定四边相等；平行四边形对边相等；菱形四条边相等。",
        ),
    ])

    processed = DocxProcessor.apply_custom_processing(
        docx_bytes,
        auto_layout_tables=True,
    )

    widths = _table_grid_widths(processed)
    assert widths[5] >= 2200
    assert widths[5] >= widths[3]


def test_auto_layout_gives_long_note_column_space_in_very_wide_tables():
    docx_bytes = _make_docx_table([
        ("日期", "渠道", "曝光", "点击", "CTR", "注册", "CVR", "付费", "ARPU", "GMV", "ROI", "备注"),
        ("2026-06-01", "Search", "120,000", "8,420", "7.02%", "1,260", "14.96%", "168", "39.8", "6,686.4", "1.82", "搜索渠道流量稳定，点击率较上周略有提升。"),
        ("2026-06-02", "Feed", "96,500", "5,930", "6.15%", "842", "14.20%", "101", "36.1", "3,646.1", "1.34", "信息流素材更换后点击率提升，但付费转化仍偏低。"),
        ("2026-06-03", "Referral", "18,240", "2,108", "11.56%", "690", "32.73%", "122", "42.7", "5,209.4", "3.91", "推荐渠道量级较小但质量较高。"),
        ("2026-06-04", "Email", "42,300", "3,012", "7.12%", "510", "16.93%", "74", "31.5", "2,331.0", "1.76", "邮件召回用户对促销活动响应明显。"),
    ])

    processed = DocxProcessor.apply_custom_processing(
        docx_bytes,
        auto_layout_tables=True,
    )

    widths = _table_grid_widths(processed)
    assert widths[-1] >= 1600
    assert widths[-1] > widths[0]
    assert all(width > 0 for width in widths)
